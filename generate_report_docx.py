import os
import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx is not installed. Attempting to install it...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
    except Exception as e:
        print(f"Failed to install python-docx: {e}")
        print("Please install it manually using: pip install python-docx")
        sys.exit(1)

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding for table cells in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcMar.append(node)
    tcPr.append(tcMar)

def create_element(name):
    return OxmlElement(name)

def set_table_borders(table):
    """Apply professional thin borders to table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def build_report():
    doc = docx.Document()
    
    # Page setup - Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    
    # Configure Normal Style (Body Text)
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Helper to add headings with custom spacing and keep_with_next
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12.5)
        return p

    # --- COVER PAGE ---
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("SUMMER INTERNSHIP PROJECT REPORT")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.name = 'Times New Roman'
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run(
        "\nOI CF – Driving Growth through Top Outlet, SAMT Outlet & Range Selling (GT-Sales) "
        "and Route-to-Market Optimization"
    )
    run_sub.bold = True
    run_sub.font.size = Pt(14)
    run_sub.font.name = 'Times New Roman'
    
    for _ in range(2):
        doc.add_paragraph()
        
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_org = p_org.add_run("ORGANIZATION:\nTata Consumer Products Limited (TCPL)")
    run_org.bold = True
    run_org.font.size = Pt(13)
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_by = p_by.add_run(
        "Submitted By:\n"
        "Niket Parikh\n"
        "Roll No: IPM04054\n"
        "Class: Integrated Programme in Management (IPM), Year 5\n"
    )
    run_by.font.size = Pt(12)
    
    p_to = doc.add_paragraph()
    p_to.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_to = p_to.add_run(
        "Submitted To:\n"
        "Academic Committee\n"
        "Indian Institute of Management, Rohtak (IIM Rohtak)\n\n"
        "In Partial Fulfillment of the Requirements for the Integrated Programme in Management\n\n"
        "Date of Submission: June 26, 2026"
    )
    run_to.font.size = Pt(11)
    run_to.italic = True
    
    doc.add_page_break()
    
    # --- TABLE OF CONTENTS ---
    add_heading_1("TABLE OF CONTENTS")
    toc_items = [
        ("1. Executive Summary", 3),
        ("2. Introduction", 4),
        ("   2.1 Background of the Organization", 4),
        ("   2.2 Objective of the Internship", 4),
        ("   2.3 Scope of the Internship", 4),
        ("   2.4 Methodology Used", 5),
        ("3. Description of the Internship", 6),
        ("   3.1 Overview of the Department/Division", 6),
        ("   3.2 Roles and Responsibilities", 6),
        ("   3.3 Tasks and Projects Undertaken", 7),
        ("4. Learning and Experiences", 8),
        ("   4.1 Skills Developed", 8),
        ("   4.2 Knowledge Acquired", 8),
        ("   4.3 Challenges Faced", 9),
        ("   4.4 Lessons Learned", 10),
        ("5. Analysis and Observations", 11),
        ("   5.1 Industry Analysis", 11),
        ("   5.2 Organizational Analysis", 11),
        ("   5.3 Process Analysis", 12),
        ("   5.4 Performance Analysis", 12),
        ("6. Contribution and Achievements", 14),
        ("   6.1 Accomplishments", 14),
        ("   6.2 Projects Completed", 14),
        ("   6.3 Initiatives Taken", 14),
        ("   6.4 Results Achieved", 15),
        ("7. Feedback and Evaluation", 16),
        ("   7.1 Feedback from Supervisor/Mentor", 16),
        ("   7.2 Self-Evaluation", 16),
        ("8. Conclusion", 17),
        ("9. Recommendations", 18),
        ("10. Appendix", 19),
        ("    10.1 Summary of Outlet Interventions by Tier", 19),
        ("    10.2 The R.O.T.A.T.E Range-Selling SOP", 20),
        ("    10.3 Beat Correction Data for Vaikunth DB", 21)
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        dots_count = 80 - len(item)
        dots = "." * max(dots_count, 5)
        p.add_run(f"{item} {dots} {page}")
        
    doc.add_page_break()
    
    # --- 1. EXECUTIVE SUMMARY ---
    add_heading_1("1. EXECUTIVE SUMMARY")
    doc.add_paragraph(
        "When I stepped into Tata Consumer Products Limited (TCPL) for my summer internship, my goal was clear: "
        "find out why the recently acquired Capital Foods (CF) and Organic India (OI) brands weren't moving as "
        "expected in General Trade (GT) stores across the South Gujarat and Ahmedabad markets, and fix it. What I "
        "quickly realized on the ground was that consumer demand wasn't the issue. Instead, the bottleneck lay "
        "in the daily, messy realities of last-mile execution. Distributor Sales Representatives (DSRs) were running "
        "on overlapping, inefficient beat routes, store owners were frustrated with delayed display payouts, and "
        "salesmen were hesitant to pitch premium extensions out of fear that they would just sit on shelves as dead stock."
    )
    doc.add_paragraph(
        "To address these challenges, I worked on two core areas. First, I conducted a detailed diagnostic across "
        "150+ retail outlets. By dividing these outlets into Type A, B, and C categories, I customized our pitch "
        "and display strategies for each segment. I developed the R.O.T.A.T.E framework—a step-by-step field guide "
        "that shows salesmen how to introduce new products using low-risk trials. During my pilot runs, this "
        "approach led to concrete wins, such as generating a single-visit order worth ₹30,000–35,000 at Falguni Gruh Udyog "
        "and reactivating several accounts that had stopped ordering from us due to past servicing issues."
    )
    doc.add_paragraph(
        "Second, I tackled route-to-market logistics at Vaikunth DB. The distributor's salesmen were stepping on "
        "each other's toes in the market, which left several high-potential areas under-serviced. I physically "
        "audited their routes, mapped the coordinates of the stores, and re-designed their territories into 42 "
        "new theoretical beats. By setting up a daily WhatsApp reporting group, I helped improve DSR compliance to "
        "84%. This effort gave us a clean database for route mapping and resolved overlapping coverage."
    )
    doc.add_paragraph(
        "Ultimately, this project showed me that in the FMCG sector, robust distribution and strong distributor "
        "relationships are what actually turn brand potential into sales."
    )
    
    doc.add_page_break()
    
    # --- 2. INTRODUCTION ---
    add_heading_1("2. INTRODUCTION")
    
    add_heading_2("2.1 Background of the Organization")
    doc.add_paragraph(
        "Tata Consumer Products Limited (TCPL) is the focused fast-moving consumer goods (FMCG) arm of the Tata Group. "
        "The company has a strong global footprint with major brands like Tata Tea, Tetley, Tata Salt, and Tata Sampann. "
        "In early 2024, TCPL acquired Capital Foods (known for Ching's Secret and Smith & Jones) and Organic India. This "
        "acquisition was a strategic move to enter the fast-growing premium wellness and convenience food segments."
    )
    doc.add_paragraph(
        "However, integrating these new brands into TCPL's vast distribution network is a complex task. Ching's Secret "
        "is a high-volume, mass-market brand, whereas Organic India is a premium, lower-velocity range. Finding a way to "
        "sell both effectively through the same General Trade (GT) and Stand-Alone Modern Trade (SAMT) channels is one "
        "of the company's key sales priorities today."
    )
    
    add_heading_2("2.2 Objective of the Internship")
    doc.add_paragraph(
        "My internship was structured around five core goals:\n"
        "• Run a Field Diagnostic: Analyze the current state of Outlet Initiatives, Coverage & Frequency (OI CF) to identify where our sales execution was falling short.\n"
        "• Profile Retail Outlets: Map and study top outlets and SAMTs to design targeted sales pitches.\n"
        "• Create a Range-Selling Guide: Develop a simple, practical Standard Operating Procedure (SOP) that salesmen can use to sell a wider range of products.\n"
        "• Optimize Beat Routes: Audit and re-map the beat routes of Vaikunth DB's salesmen to make their daily market coverage more efficient.\n"
        "• Run Field Pilots: Implement these strategies in the market during my internship to test if they actually work."
    )
    
    add_heading_2("2.3 Scope of the Internship")
    doc.add_paragraph(
        "My field research and pilot runs were focused on the GT and SAMT channels within the South Gujarat / West 1 region, "
        "specifically covering the Ahmedabad and Vadodara urban markets. I worked closely with Vaikunth DB and Darsh DB, "
        "auditing a sample of 150+ priority retail outlets across categories like green teas, packaged noodles, chutneys, and spices."
    )
    
    add_heading_2("2.4 Methodology Used")
    doc.add_paragraph(
        "I followed a hands-on, action-oriented approach in the field:\n"
        "1. Initial Research: I started by studying TCPL’s internal catalogs, sales targets, and the Mavic app’s user guide.\n"
        "2. Field Visits: I spent weeks traveling with DSRs on their daily runs, counting stock on shelves, correcting displays, and talking to store owners about their concerns.\n"
        "3. Retailer Classification: Based on what I observed, I grouped the 150 outlets into Type A (large/premium), Type B (medium/growth), and Type C (small/standard) stores.\n"
        "4. Testing Interventions: I ran pilot tests in 15 selected stores, focusing on manual stock counting, shelf display corrections, and pitching small trial orders.\n"
        "5. Beat Optimization: I gathered GPS coordinates and store lists from the 7 DSRs at Vaikunth DB, resolved routing overlaps, and mapped out 42 new beat routes using Excel."
    )
    
    doc.add_page_break()
    
    # --- 3. DESCRIPTION OF THE INTERNSHIP ---
    add_heading_1("3. DESCRIPTION OF THE INTERNSHIP")
    
    add_heading_2("3.1 Overview of the Department/Division")
    doc.add_paragraph(
        "I was placed in the General Trade (GT) Sales division of TCPL. GT is still the primary channel for FMCG sales in India, "
        "relying on a network of distributors and local kirana stores. The division operates under a clear hierarchy to manage "
        "these distributors. The direct line of reporting went from Territory Sales Executive (TSE) to Area Sales Manager "
        "(ASM - Shahrukh Ahmed), Regional Sales Manager (RSM - Sushant Jangale), and Channel Head West (CH - Punit Kadam)."
    )
    
    add_heading_2("3.2 Roles and Responsibilities")
    doc.add_paragraph(
        "As a Sales & Marketing Management Trainee, my responsibilities included:\n"
        "• Auditing the daily market visits of DSRs under Vaikunth DB and Darsh DB.\n"
        "• Identifying and resolving delivery and payment issues between the distributor and retailers.\n"
        "• Helping salesmen improve their sales pitches for new launches like the Organic India wellness range and Ching's Korean Ramen.\n"
        "• Mapping out beat routes and building a database of verified retail outlets."
    )
    
    add_heading_2("3.3 Tasks and Projects Undertaken")
    doc.add_paragraph(
        "My work was divided into two main projects:\n\n"
        "Project 1: Driving Range Selling and Visibility\n"
        "I focused on expanding the number of SKUs carried by retail outlets. By auditing 150 stores, I realized that DSRs "
        "rarely pitched new products because they were afraid the stock wouldn't sell. To overcome this, I developed and "
        "tested the R.O.T.A.T.E framework, which gives DSRs a step-by-step method to pitch low-risk trial orders to store owners.\n\n"
        "Project 2: Restructuring Beat Routes at Vaikunth DB\n"
        "I addressed routing inefficiencies. The 7 DSRs at Vaikunth DB had overlapping territories, which meant some "
        "markets were visited twice while others were completely ignored. I audited their routes, gathered clean outlet "
        "data, and designed 42 new theoretical beats. I also set up a daily WhatsApp reporting group to track DSR compliance."
    )
    
    doc.add_page_break()
    
    # --- 4. LEARNING AND EXPERIENCES ---
    add_heading_1("4. LEARNING AND EXPERIENCES")
    
    add_heading_2("4.1 Skills Developed")
    doc.add_paragraph(
        "• Leading Without Authority: I had to convince DSRs and distributors to change their daily routines and adopt new processes, even though they did not report to me directly.\n"
        "• Route Optimization: I learned how to analyze geographic data to design beat routes that minimize travel time and maximize the number of store visits.\n"
        "• Handling B2B Objections: I spent time on the field negotiating with store owners who were unhappy about delayed display payouts or slow deliveries.\n"
        "• Creating Field SOPs: I learned to design simple, practical processes (like the R.O.T.A.T.E framework) that salesmen can easily follow during busy market runs."
    )
    
    add_heading_2("4.2 Knowledge Acquired")
    doc.add_paragraph(
        "• How GT Distribution Works: I got a close look at credit cycles, distributor margins, and how inventory moves from warehouses to retail shelves.\n"
        "• Niche vs. Mass Brand Management: I saw first-hand the challenges of selling a premium brand (Organic India) alongside a mass-market brand (Capital Foods) through the same sales force.\n"
        "• Sales Automation Tools: I analyzed how TCPL uses SFA software (the Mavic app) and identified where the database needed cleaner, more accurate inputs."
    )
    
    add_heading_2("4.3 Challenges Faced")
    doc.add_paragraph(
        "• Uncooperative DSRs: Initially, salesmen were hesitant to share their coordinates or report their daily visits because they felt they were being micromanaged.\n"
        "• Inaccurate App Data: The Mavic app’s retail database was outdated, forcing me to collect store details and coordinates manually.\n"
        "• Logistical Issues: Vaikunth DB was located far from its primary markets, which caused delivery delays and frustrated retailers.\n"
        "• High Salesman Attrition: Frequent turnover among Darsh DB's DSRs made it difficult to maintain consistent servicing.\n"
        "• Lack of ID Card: Not having an official company ID card during my first few weeks made some retailers suspicious of my questions.\n"
        "• Strong Competition: Competitors like Keya and Wok Tok visited stores regularly and frequently pushed our products off the primary shelves."
    )
    
    add_heading_2("4.4 Lessons Learned")
    doc.add_paragraph(
        "• Execution Wins on the Ground: A brilliant marketing strategy is useless if the distributor fails to deliver the stock or if the salesman skips his weekly visit.\n"
        "• Start Small to Build Trust: Kirana owners hate taking risks. Pitching a small trial order (like 3 packets) is much easier than forcing them to buy a whole case, and it helps build long-term trust.\n"
        "• Servicing Consistency is Key: Store owners are much more willing to give us prime shelf space if they know our salesman will show show up reliably every single week."
    )
    
    doc.add_page_break()
    
    # --- 5. ANALYSIS AND OBSERVATIONS ---
    add_heading_1("5. ANALYSIS AND OBSERVATIONS")
    
    add_heading_2("5.1 Industry Analysis")
    doc.add_paragraph(
        "The Indian FMCG market is moving in two distinct directions. On one hand, young urban consumers are demanding quick "
        "convenience foods like instant noodles and cooking sauces. On the other hand, there is a growing demand for health "
        "and wellness products like organic honey, green teas, and natural sweeteners. In GT stores, shelf space is extremely "
        "limited, and brands have to fight constantly to keep their products visible."
    )
    
    add_heading_2("5.2 Organizational Analysis")
    doc.add_paragraph(
        "TCPL's acquisition of Capital Foods and Organic India aligns well with these market trends. However, selling a "
        "fast-moving, low-margin noodle brand alongside a slow-moving, high-margin wellness tea brand using the same sales team "
        "is challenging. DSRs naturally focus on high-velocity items like Ching's Secret and tend to neglect premium Organic India "
        "products because they require a different, more patient sales pitch."
    )
    
    add_heading_2("5.3 Process Analysis")
    doc.add_paragraph(
        "During my field audits, I noticed that DSRs rarely followed the company's standard sales call model. Instead of checking "
        "shelf inventory or pitching new launches, DSRs would simply ask the retailer 'Kya chahiye?' (What do you need?), "
        "write down a quick order, and leave. This hurried approach meant we were missing out on potential sales and failing to "
        "introduce new products."
    )
    
    add_heading_2("5.4 Performance Analysis")
    doc.add_paragraph(
        "My diagnostics highlighted several execution gaps:\n"
        "• Gaps in Product Range: Many high-potential stores that sold a lot of Ching's sauces carried no noodles at all.\n"
        "• Poor Display Placement: Organic India products were often placed on bottom shelves or mixed with general foods, instead of being displayed in the premium tea section.\n"
        "• Inaccurate App Reporting: DSRs would sometimes mark stores as 'closed' or log a 'zero sale' on the Mavic app without actually visiting, just to meet their daily call targets.\n"
        "• Payout Delays: Delayed display payments (like at Hari Om Super Market) made retailers angry, leading them to hide our stock or refuse to carry new ranges."
    )
    
    doc.add_page_break()
    
    # --- 6. CONTRIBUTION AND ACHIEVEMENTS ---
    add_heading_1("6. CONTRIBUTION AND ACHIEVEMENTS")
    
    add_heading_2("6.1 Accomplishments")
    doc.add_paragraph(
        "• Conducted field audits and diagnosed sales execution issues across 150+ retail outlets.\n"
        "• Analyzed the distribution operations of Vaikunth DB and Darsh DB.\n"
        "• Set up a structured communication loop that helped DSRs and distributors coordinate better."
    )
    
    add_heading_2("6.2 Projects Completed")
    doc.add_paragraph(
        "• Designed 42 new theoretical beat routes to eliminate overlapping coverage for Vaikunth DB.\n"
        "• Built a Retail Master Sheet containing verified names, phone numbers, addresses, and coordinates for Vaikunth DB's outlets.\n"
        "• Created the R.O.T.A.T.E framework as a simple field guide for range selling."
    )
    
    add_heading_2("6.3 Initiatives Taken")
    doc.add_paragraph(
        "• Launched a WhatsApp daily reporting group for Vaikunth DB's salesmen, which reached an 84% compliance rate.\n"
        "• Conducted manual inventory counts and display corrections during my joint field runs to show DSRs how to identify sales opportunities."
    )
    
    add_heading_2("6.4 Results Achieved")
    doc.add_paragraph(
        "• Introduced New Ranges: Successfully seeded new product categories in 10+ pilot outlets.\n"
        "• Improved Shelf Visibility: Secured better shelf space and block displays in 40+ stores.\n"
        "• Reactivated Dormant Stores: Resumed servicing for 4 stores that had stopped buying from TCPL due to unresolved delivery disputes.\n"
        "• Generated High-Value Orders: Personally secured a ₹30,000–35,000 order at Falguni Gruh Udyog by doing a manual stock count and organizing their display.\n"
        "• Cleaned Route Data: Gathered verified coordinate data for Vaikunth DB's outlets, providing a clean database for future beat mapping on the Mavic app."
    )
    
    doc.add_page_break()
    
    # --- 7. FEEDBACK AND EVALUATION ---
    add_heading_1("7. FEEDBACK AND EVALUATION")
    
    add_heading_2("7.1 Feedback from Supervisor/Mentor")
    doc.add_paragraph(
        "My Area Sales Manager, Shahrukh Ahmed, noted that the beat correction project provided clean, practical field data that "
        "resolved several routing disputes among DSRs. He also highlighted that the pilot results at Falguni Gruh Udyog and "
        "Bajrang Super Market proved how much we can increase order values by doing active stock counts. The regional skip managers "
        "(Sushant Jangale and Punit Kadam) advised that we should focus on training DSRs on the R.O.T.A.T.E framework to make "
        "these sales improvements sustainable."
    )
    
    add_heading_2("7.2 Self-Evaluation")
    doc.add_paragraph(
        "This internship was a great lesson in the operational challenges of FMCG sales. I succeeded in building a good working "
        "relationship with the field team and resolving routing issues. However, I also realized that changing the behavior of "
        "DSRs who are used to quick, transactional sales calls is a slow process that requires constant reinforcement and close monitoring."
    )
    
    doc.add_page_break()
    
    # --- 8. CONCLUSION ---
    add_heading_1("8. CONCLUSION")
    doc.add_paragraph(
        "My 2-month internship at TCPL focused on solving last-mile sales execution and distribution issues in the GT channel. "
        "The diagnostic audits confirmed that while there is strong demand for our brands, our sales growth was being limited "
        "by overlapping beats and inconsistent market visits."
    )
    doc.add_paragraph(
        "By re-mapping Vaikunth DB's territory into 42 distinct beats and building a verified Retail Master Sheet, I helped create "
        "a cleaner route-to-market database. Furthermore, our pilot runs of the R.O.T.A.T.E range-selling framework showed that "
        "using low-risk trials and performing active stock counts can significantly increase order sizes, offering a clear way to "
        "improve sales execution across the region."
    )
    
    doc.add_page_break()
    
    # --- 9. RECOMMENDATIONS ---
    add_heading_1("9. RECOMMENDATIONS")
    doc.add_paragraph(
        "Based on my observations, I recommend the following steps:\n"
        "1. Use a Smart Reorder System: Set up an automated inventory check for top-selling products (like Ching's Schezwan Chutney) to prevent stockouts in high-volume stores.\n"
        "2. Ensure Timely Display Payouts: Link display compliance with regular payouts, and make sure retailers receive these incentives on time to rebuild their trust.\n"
        "3. Equip DSRs with Better Sales Tools: Provide DSRs with physical or digital catalogs that show new launches, replacing outdated pricing sheets.\n"
        "4. Audit Distributor Deliveries: Regularly check distributor delivery schedules, especially for stores located far from the warehouse, to prevent delivery delays."
    )
    
    doc.add_page_break()
    
    # --- 10. APPENDIX ---
    add_heading_1("10. APPENDIX")
    
    add_heading_2("10.1 Summary of Outlet Interventions by Tier")
    
    # Styled table
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    headers = ["Outlet Tier", "Target Profile", "Key Issues Observed", "Core Sales Strategy", "Pilot Evidence"]
    
    # Style Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F497D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            
    # Table Content Data
    table_data = [
        ("Type A (Premium)", "Falguni Gruh Udyog\nBajrang Super Market\nD2D Mart", "Incomplete range\nPassive displays", "Push larger order size\nCorrect shelf visibility\nIntroduce premium SKUs", "Falguni: ₹30-35k order after manual count\nBajrang: Expanded OI range"),
        ("Type B (Growth)", "Madhur Super Market\nSaavi Gruh Udyog\nMohit Dept Store", "Delivery/payment disputes\nInconsistent visits", "Re-establish weekly visits\nPitch proven extensions", "Saavi: Resolved dispute, resumed deliveries, ordered Ramen & green tea"),
        ("Type C (Small)", "Vraj Super Market\nDharmi Super Market\nRiddhi Siddhi SM", "Expired shelf stock\nRetailer distrust", "Re-introduce core SKUs\nUse 3-unit trial orders", "Vraj: Re-introduced OI teas\nRiddhi Siddhi: Ordered Pink Salt trial")
    ]
    
    for row_idx, data in enumerate(table_data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
                
    doc.add_paragraph()
    
    add_heading_2("10.2 The R.O.T.A.T.E Range-Selling SOP")
    doc.add_paragraph(
        "To expand our product range in retail outlets, DSRs should follow these six steps:\n"
        "• R - Range Audit: Count the active SKUs on the shelves and check for expired stock.\n"
        "• O - Opportunity Pitch: Identify missing categories (like soups or pasta masalas) that are selling well in nearby stores.\n"
        "• T - Trial Offer: Offer a small trial quantity (e.g., 3 units) instead of forcing a full case to reduce the retailer's risk.\n"
        "• A - Assortment Display: Set up a clean shelf block where products are grouped together by brand.\n"
        "• T - Trust & Schemes: Explain active retailer schemes and payouts clearly, and note any unpaid incentives.\n"
        "• E - Execution & Follow-up: Schedule a follow-up visit for the same day the following week to track sales and take reorders."
    )
    
    add_heading_2("10.3 Beat Correction Data for Vaikunth DB")
    doc.add_paragraph(
        "• Original State: 7 DSRs operated with overlapping territories, resulting in inefficient routing and irregular outlet visits.\n"
        "• Corrected State: Operational areas were re-demarcated into 42 distinct theoretical beats.\n"
        "• Compliance Tool: Set up a WhatsApp reporting group for salesmen, requiring daily check-ins in the following format:\n"
        "  - [Outlet Name]\n"
        "  - [Contact Person]\n"
        "  - [Contact Number]\n"
        "  - [Beat Name]\n"
        "  - [Address / Location Pin]\n"
        "• Result: Achieved 84% daily reporting compliance, providing verified coordinate data for the Retail Master Sheet and Mavic app integration."
    )
    
    # Save document
    out_file = '/Users/ketanparikh/Desktop/Antigravity Work/SIP_Report_Niket_Parikh.docx'
    doc.save(out_file)
    print(f"Report saved successfully to: {out_file}")

if __name__ == '__main__':
    build_report()
